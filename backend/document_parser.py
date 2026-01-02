import os
import re
from PyPDF2 import PdfReader
from docx import Document
from backend.utils import setup_logger

logger = setup_logger(__name__)


def parse_document(file_path):
    """
    Main dispatcher for parsing documents based on file type.

    Args:
        file_path: Path to document

    Returns:
        dict: {
            "success": bool,
            "text": str,
            "error": str or None,
            "char_count": int
        }
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    try:
        if ext == ".pdf":
            text = extract_text_from_pdf(file_path)
        elif ext == ".docx":
            text = extract_text_from_docx(file_path)
        elif ext == ".txt":
            text = extract_text_from_txt(file_path)
        else:
            return {
                "success": False,
                "text": "",
                "error": f"Unsupported file type: {ext}",
                "char_count": 0
            }

        # Clean the extracted text
        cleaned_text = clean_text(text)

        return {
            "success": True,
            "text": cleaned_text,
            "error": None,
            "char_count": len(cleaned_text)
        }

    except Exception as e:
        logger.error(f"Error parsing {file_path}: {str(e)}")
        return {
            "success": False,
            "text": "",
            "error": str(e),
            "char_count": 0
        }


def extract_text_from_pdf(file_path):
    """
    Extract text from PDF using PyPDF2.

    Args:
        file_path: Path to PDF file

    Returns:
        str: Extracted text
    """
    try:
        reader = PdfReader(file_path)
        text = ""

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        if not text.strip():
            logger.warning(f"No text extracted from PDF: {file_path}")

        return text

    except Exception as e:
        logger.error(f"Error reading PDF {file_path}: {str(e)}")
        raise


def extract_text_from_docx(file_path):
    """
    Extract text from DOCX using python-docx.

    Args:
        file_path: Path to DOCX file

    Returns:
        str: Extracted text
    """
    try:
        doc = Document(file_path)
        text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        if not text.strip():
            logger.warning(f"No text extracted from DOCX: {file_path}")

        return text

    except Exception as e:
        logger.error(f"Error reading DOCX {file_path}: {str(e)}")
        raise


def extract_text_from_txt(file_path):
    """
    Extract text from plain text file.

    Args:
        file_path: Path to TXT file

    Returns:
        str: File contents
    """
    try:
        # Try UTF-8 first, fall back to latin-1 if that fails
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    return text
            except UnicodeDecodeError:
                continue

        # If all encodings fail, read as binary and decode with errors='ignore'
        with open(file_path, 'rb') as f:
            text = f.read().decode('utf-8', errors='ignore')
            return text

    except Exception as e:
        logger.error(f"Error reading TXT {file_path}: {str(e)}")
        raise


def clean_text(text):
    """
    Clean and normalize extracted text.

    Args:
        text: Raw extracted text

    Returns:
        str: Cleaned text
    """
    if not text:
        return ""

    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)

    # Remove excessive newlines (more than 2 consecutive)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Strip leading/trailing whitespace
    text = text.strip()

    return text
